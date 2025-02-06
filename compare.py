import os
import math
import pandas as pd
from typing import List, Dict, Union
from pydantic import BaseModel
from openai import AzureOpenAI, OpenAI
from dotenv import load_dotenv
import streamlit as st
import logging  # ログ機能の追加
from datetime import datetime  # datetimeモジュールを追加
import PyPDF2  # PDFファイル用ライブラリ
import docx     # Wordファイル用ライブラリ
import io       # ファイルオブジェクトを扱うため

def read_file_with_encoding(file_content: bytes) -> str:
    """
    複数のエンコーディングを試してファイルを読み込む関数
    """
    encodings = ['utf-8', 'shift_jis', 'cp932', 'euc_jp']
    
    for encoding in encodings:
        try:
            text = file_content.decode(encoding)
            if text.strip():  # 空でない場合のみ成功とみなす
                logging.info(f"{encoding}でのデコードに成功しました。")
                return text
            logging.warning(f"{encoding}でデコードは成功しましたが、テキストが空です。")
        except UnicodeDecodeError:
            logging.warning(f"{encoding}でのデコードに失敗しました。")
            continue
    
    raise ValueError("すべてのエンコーディングでデコードに失敗しました。")

def read_file_content(uploaded_file) -> bytes:
    """
    アップロードされたファイルの内容を読み込む関数
    """
    try:
        # ファイルポインタを先頭に戻す
        uploaded_file.seek(0)
        content = uploaded_file.getvalue()  # read()の代わりにgetvalue()を使用
        return content
    except Exception as e:
        logging.error(f"ファイル読み込みエラー: {e}")
        raise

def read_word_chunks(uploaded_file) -> List[str]:
    """
    Wordファイルからテキストを抽出し、段落ごとにチャンク化する関数。
    """
    text_chunks = []
    try:
        content = read_file_content(uploaded_file)
        doc_buffer = io.BytesIO(content)
        doc = docx.Document(doc_buffer)
        
        # テーブルのテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_chunks.append(cell.text.strip())
        
        # 段落のテキストを抽出
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 空の段落を除外
                text_chunks.append(paragraph.text)
        
        # 空のチャンクリストをチェック
        if not text_chunks:
            logging.warning("Wordファイルから抽出されたテキストが空です。")
            return []
            
        return text_chunks
    except Exception as e:
        logging.error(f"Wordファイル読み込みエラー: {e}")
        st.error(f"Wordファイルの読み込みに失敗しました: {e}")
        return []

def read_pdf_chunks(uploaded_file) -> List[str]:
    """
    PDFファイルからテキストを抽出し、ページごとにチャンク化する関数。
    """
    text_chunks = []
    try:
        content = read_file_content(uploaded_file)
        pdf_buffer = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_buffer)
        
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():  # 空のページを除外
                # 長い文章を適切な長さで分割
                paragraphs = text.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        text_chunks.append(paragraph.strip())
        
        # 空のチャンクリストをチェック
        if not text_chunks:
            logging.warning("PDFファイルから抽出されたテキストが空です。")
            return []
            
        return text_chunks
    except Exception as e:
        logging.error(f"PDFファイル読み込みエラー: {e}")
        st.error(f"PDFファイルの読み込みに失敗しました: {e}")
        return []

# .envファイルから環境変数を読み込む
load_dotenv()

# ログ設定
logging.basicConfig(
    filename='compare_tool.log',
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    encoding='utf-8'
)
logging.getLogger('dotenv').setLevel(logging.WARNING)  # dotenvのログレベルを下げる

# API設定の取得
API_TYPE = os.getenv("API_TYPE", "azure")  # デフォルトはazure
if API_TYPE.lower() == "azure":
    client = AzureOpenAI(
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
    )
    # Azure OpenAIのモデルデプロイメント名
    EMBEDDING_MODEL = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
    CHAT_MODEL = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT")
else:
    client = OpenAI(
        api_key=os.getenv("OPENAI_API_KEY")
    )
    # OpenAIのモデル名
    EMBEDDING_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
    CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4-turbo-preview")

# Pydanticモデル：GPT-4oの応答を構造化する
class CompareOutput(BaseModel):
    is_similar: bool
    difference_summary: str

def get_embedding(text: str) -> List[float]:
    """
    テキストを Embedding ベクトルに変換する関数。
    """
    if API_TYPE.lower() == "azure":
        response = client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=text
        )
    else:
        response = client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=text
        )
    embedding = response.data[0].embedding
    return embedding

def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """
    コサイン類似度を計算する。
    """
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = math.sqrt(sum(a * a for a in vec1))
    norm2 = math.sqrt(sum(b * b for b in vec2))
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return dot_product / (norm1 * norm2)

def compare_texts_with_gpt4o(base_text: str, compare_texts: Union[str, List[str]]) -> CompareOutput:
    """
    GPT-4o を使って 本社規定テキスト と 子会社規定テキスト/テキスト群 を比較し、
    似ているかどうか (is_similar) と差分サマリ (difference_summary) を
    構造化データとして返す。
    子会社規定が複数のチャンクに分割されている場合も対応。
    """
    messages = [
        {
            "role": "system",
            "content": (
                "あなたは法務の専門家アシスタントです。"
                "出力はJSON形式で、is_similarとdifference_summaryの2項目を必ず返してください。"
                "これから提示する【本社規定項目】と【子会社規定】を比較し、"
                "子会社規定が本社規定の内容を実質的に満たしているか判断してください。"
                "もし内容が実質的に同じ(本社規定を満たしている)と判断できる場合はis_similar=Trueにしてください。"
                "一部でも本社規定を満たしていない場合はis_similar=Falseとしてください。"
                "重要な注意点として、子会社規定にのみ存在する記述（例えば、適用範囲の詳細など）は、本社規定の充足性判断には不要です。"
                "これらの子会社規定にしかない情報は、差分サマリに含めないでください。"
                "比較対象は、本社規定の項目が子会社規定に実質的に含まれているかどうかのみです。"
                "子会社規定が本社規定の意図を汲み、同等の内容を別の表現で記述している場合でも、類似していると判断してください。"
                "子会社規定は複数のチャンクに分割されている可能性があります。すべてのチャンクを考慮して、全体として本社規定を満たしているか判断してください。"
                "差分サマリは、本社規定を満たしていない部分のみを具体的に記述してください。"
                "もし差分がない場合は、difference_summaryは空にしてください。"
            )
        },
        {
            "role": "user",
            "content": (
                "以下の【本社規定項目】の内容を【子会社規定】が満たしているか比較してください。\n"
                "【子会社規定】は複数のテキストチャンクで構成されている場合があります。\n"
                "すべてのチャンクを考慮して、全体として本社規定を満たしているか判断してください。\n"
                "もし文面が違う場合は、その違いを差分サマリとして簡潔に教えてください。\n"
                "出力はJSON形式で日本語で返してください。\n\n"
                "【本社規定項目】\n"
                f"{base_text}\n\n"
                "【子会社規定】\n"
                f"{chr(10).join(compare_texts) if isinstance(compare_texts, list) else compare_texts}\n"
            )
        }
    ]

    completion = client.beta.chat.completions.parse(
        model=CHAT_MODEL,
        messages=messages,
        response_format=CompareOutput,
    )
    return completion.choices[0].message.parsed

def generate_markdown_report(
    checklist: List[Dict],
    subsidiary_chunks: List[Dict],
    top_k: int = 3,
    similarity_threshold: float = 0.5
) -> str:
    """
    チェックリスト (本社規定の項目一覧) と 子会社規定のチャンク を比較して、
    Embeddings で候補を絞りつつ GPT-4o で差分要約。
    結果を Markdown で返す。

    改善点:
    1. 閾値を超えたチャンクはすべて関連チャンクとみなし、GPT-4oに問い合わせる。
    2. 1件も閾値以上がなければ「本社規定のみ」として扱う。
    3. GPT-4o応答(is_similar)がTrueのものが1つでもあれば「該当あり」とみなす。
    """

    logging.info("レポート生成開始")

    # 子会社チャンクに embedding が無い場合はここで埋め込む
    for chunk in subsidiary_chunks:
        if "embedding" not in chunk:
            chunk["embedding"] = get_embedding(chunk["text"])

    md_report = []
    md_report.append("# チェックリスト比較レポート\n")

    for item in checklist:
        item_id = item["id"]
        item_text = item["text"]

        logging.info(f"チェック項目ID: {item_id} の比較開始")
        logging.debug(f"本社規定項目テキスト: {item_text}")

        # 本社規定アイテムの embedding
        item_embedding = get_embedding(item_text)

        # 子会社規定チャンクとの類似度を計算 → 全件スコアリング
        scored_chunks = []
        for chunk in subsidiary_chunks:
            sim = cosine_similarity(item_embedding, chunk["embedding"])
            scored_chunks.append({
                "id": chunk["id"],
                "text": chunk["text"],
                "similarity": sim
            })

        # 類似度でソート
        scored_chunks.sort(key=lambda x: x["similarity"], reverse=True)

        # (1) 閾値以上のチャンクだけをリストアップ
        related_chunks = [c for c in scored_chunks if c["similarity"] >= similarity_threshold]
        # 閾値を超えるものがなければ、「本社規定にのみ存在」として扱う
        if not related_chunks:
            md_report.append(f"## チェック項目: {item_id}")
            md_report.append(f"**本社規定項目テキスト**: {item_text}\n")
            md_report.append("- 対応チャンク: **該当なし**（本社規定にのみ存在）\n")
            md_report.append("---")
            logging.info(f"チェック項目ID: {item_id} は関連チャンクなしと判定。")
            continue

        # (2) 閾値以上のチャンクがある場合、GPT-4o で順次判定
        found_fulfilled = False   # 一つでも is_similar = True があったか
        chunk_results = []        # 全チャンクの結果を保持

        # 上位N件に限定したい場合はここで[:top_k]にする
        related_chunk_texts = [c["text"] for c in related_chunks] # 関連チャンクのテキストをリストで取得
        if related_chunk_texts: # 関連チャンクが1つ以上存在する場合のみGPT-4o比較
            logging.debug(f"GPT-4o比較開始: {related_chunks}")
            compare_result = compare_texts_with_gpt4o(item_text, related_chunk_texts) # 複数チャンクをまとめて比較
            chunk_results.append({
                "chunk_ids": [c["id"] for c in related_chunks], # チャンクIDをリストで保持
                "similarities": [c["similarity"] for c in related_chunks], # 類似度をリストで保持
                "is_similar": compare_result.is_similar,
                "difference_summary": compare_result.difference_summary
            })
            if compare_result.is_similar:
                found_fulfilled = True

        # Markdown出力
        md_report.append(f"## チェック項目: {item_id}")
        md_report.append(f"**本社規定項目テキスト**: {item_text}\n")

        if found_fulfilled:
            md_report.append("- <span style='color: green; font-weight: bold'>この項目は子会社規定で満たされています。</span>")
        else:
            md_report.append("- <span style='color: red; font-weight: bold'>この項目は子会社規定で不十分と判定されました。</span>")

        # 関連チャンクの比較結果一覧を表示
        md_report.append("#### 関連チャンクの判定結果 (閾値以上)")
        for cr in chunk_results:
            md_report.append(f"- チャンクID: {', '.join(cr['chunk_ids'])} (類似度: {', '.join([str(round(s, 3)) for s in cr['similarities']])})") # 複数チャンクIDと類似度を表示
            md_report.append(f"  - is_similar: {cr['is_similar']}")
            if cr["difference_summary"]:
                md_report.append(f"  - 差分: {cr['difference_summary']}")
            md_report.append("")

        md_report.append("---")

    logging.info("レポート生成完了")
    return "\n".join(md_report)

def generate_checklist_from_document(document_text: str) -> List[Dict[str, str]]:
    """
    本社規定文書からチェックリスト項目を生成する関数
    """
    messages = [
        {
            "role": "system",
            "content": (
                "あなたは法務の専門家アシスタントです。"
                "提供された規定文書から、重要な規定項目をチェックリスト形式で抽出してください。"
                "各項目は簡潔で明確な文章にしてください。"
                "出力は1行1項目の形式で、各項目は規定の重要な要件を表現するものとしてください。"
            )
        },
        {
            "role": "user",
            "content": (
                "以下の規定文書から、重要な規定項目をチェックリスト形式で抽出してください。\n\n"
                f"{document_text}"
            )
        }
    ]

    completion = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=messages,
        temperature=0
    )
    
    # 応答から項目を抽出してリスト形式に変換
    checklist_items = completion.choices[0].message.content.strip().split('\n')
    return [{"id": f"A-{i+1}", "text": item.strip()} for i, item in enumerate(checklist_items) if item.strip()]

def merge_checklists(checklists: List[List[Dict[str, str]]]) -> List[Dict[str, str]]:
    """
    複数のチェックリストを1つに統合する関数
    重複を除去し、IDを振り直します
    """
    # 全てのチェックリストアイテムのテキストを集める
    all_items = []
    seen_texts = set()  # 重複チェック用
    
    for checklist in checklists:
        for item in checklist:
            if item["text"] not in seen_texts:
                all_items.append(item["text"])
                seen_texts.add(item["text"])
    
    # 新しいIDを振り直してチェックリストを作成
    return [{"id": f"A-{i+1}", "text": text} for i, text in enumerate(all_items)]

# Streamlit アプリケーション構築
st.title("規定比較ツール")

# チェックリストを初期化
example_checklist = []

st.header("1. 本社規定")
# 本社規定のファイルアップロード
parent_uploaded_file = st.file_uploader("本社規定のファイルをアップロードしてください (PDF, Word)", type=["pdf", "docx"], key="parent_file")

st.header("2. 子会社規定チャンク")
# ファイルアップローダーを配置
uploaded_file = st.file_uploader("子会社規定のファイルをアップロードしてください (PDF, Word)", type=["pdf", "docx"])

st.sidebar.header("パラメータ設定")
top_k = st.sidebar.number_input("Top K", min_value=1, max_value=10, value=3)
similarity_threshold = st.sidebar.slider("類似度閾値", 0.0, 1.0, 0.5, step=0.1)

if st.button("比較実行"):
    logging.debug(f"【デバッグログ】比較実行ボタン押下時 example_checklist: {example_checklist}")
    
    # 本社規定ファイルの処理
    if parent_uploaded_file is not None:
        try:
            file_extension = parent_uploaded_file.name.split('.')[-1].lower()
            parent_chunks = []
            
            if file_extension == "pdf":
                parent_chunks = read_pdf_chunks(parent_uploaded_file)
                if not parent_chunks:
                    st.error("PDFからテキストを抽出できませんでした。")
                    logging.error("PDFからのテキスト抽出に失敗")
            elif file_extension == "docx":
                parent_chunks = read_word_chunks(parent_uploaded_file)
                if not parent_chunks:
                    st.error("Wordファイルからテキストを抽出できませんでした。")
                    logging.error("Wordからのテキスト抽出に失敗")
            else:
                # テキストファイルとして読み込み
                content = read_file_content(parent_uploaded_file)
                text = read_file_with_encoding(content)
                # テキストファイルの場合も適度な長さでチャンク分割
                parent_chunks = [text[i:i+4000] for i in range(0, len(text), 4000)]
            
            if not parent_chunks:
                st.error("本社規定ファイルの内容が空です。ファイルを確認してください。")
                logging.error("本社規定ファイルの内容が空です。")
            else:
                # 各チャンクに対してチェックリストを生成
                with st.spinner("チェックリストを生成中..."):
                    all_checklists = []
                    for i, chunk in enumerate(parent_chunks):
                        logging.info(f"チャンク {i+1}/{len(parent_chunks)} のチェックリスト生成中")
                        chunk_checklist = generate_checklist_from_document(chunk)
                        all_checklists.append(chunk_checklist)
                    
                    # 全チェックリストを統合
                    example_checklist = merge_checklists(all_checklists)
                    
                    if not example_checklist:
                        st.error("本社規定からチェックリストを生成できませんでした。")
                        logging.error("チェックリストの生成に失敗しました。")
                    else:
                        st.success("チェックリストを生成しました！")
                        with st.expander("チェックリストの内容を表示"):
                            for item in example_checklist:
                                st.write(f"{item['id']}: {item['text']}")
        
        except Exception as e:
            logging.error(f"ファイル読み込みエラー: {e}")
            st.error(f"ファイルの読み込みに失敗しました: {str(e)}")
    
    # 子会社規定の処理
    if not example_checklist:
        st.error("チェックリストを入力してください。")
    elif not uploaded_file:
        st.error("子会社規定ファイルをアップロードしてください。")
    else:
        with st.spinner("比較レポート作成中..."):
            example_subsidiary_chunks = []
            if uploaded_file is not None:
                file_extension = uploaded_file.name.split('.')[-1].lower()
                if file_extension == "pdf":
                    chunks_text = read_pdf_chunks(uploaded_file)
                elif file_extension == "docx":
                    chunks_text = read_word_chunks(uploaded_file)
                else:
                    st.error("PDFまたはWordファイルをアップロードしてください。")
                    chunks_text = []
                example_subsidiary_chunks = [
                    {"id": f"F-{i+1}", "text": item.strip()}
                    for i, item in enumerate(chunks_text) if item.strip()
                ]

            if example_subsidiary_chunks:
                report_md = generate_markdown_report(
                    checklist=example_checklist,
                    subsidiary_chunks=example_subsidiary_chunks,
                    top_k=top_k,
                    similarity_threshold=similarity_threshold
                )
                st.header("比較レポート")
                st.markdown(report_md, unsafe_allow_html=True)
                st.success("比較レポートが完成しました！")

                # resultフォルダを作成（存在しない場合）
                os.makedirs("result", exist_ok=True)

                # 現在日時をファイル名に含める
                now = datetime.now().strftime("%Y%m%d_%H%M%S")
                file_name = f"result/compare_report_{now}.md"

                # レポートをファイルに保存
                with open(file_name, "w", encoding="utf-8") as f:
                    f.write(report_md)
                st.info(f"レポートを {file_name} に保存しました。")
            else:
                st.error("子会社規定の処理に失敗しました。ファイルを確認してください。")

if __name__ == "__main__":
    if parent_uploaded_file is not None:
        logging.debug(f"【デバッグログ】parent_uploaded_file が None ではないと評価されました。parent_uploaded_file の種類: {type(parent_uploaded_file).__name__}")
        try:
            parent_text = parent_uploaded_file.read().decode('utf-8')
        except UnicodeDecodeError:
            logging.warning("UTF-8でのデコードに失敗しました。Shift_JISでのデコードを試みます。")
            try:
                parent_text = parent_uploaded_file.read().decode('shift_jis')
            except UnicodeDecodeError:
                logging.error("Shift_JISでのデコードにも失敗しました。")
                st.error("ファイルのデコードに失敗しました。UTF-8またはShift_JISでエンコードされたファイルをお試しください。")
                parent_text = "" # デコード失敗時は空文字を代入
            else:
                logging.info("Shift_JISでのデコードに成功しました。")
        st.write(parent_text)
        if parent_text.strip():
            logging.debug(f"【デバッグログ】parent_text.strip() が True と評価されました。parent_textの先頭50文字: {parent_text[:50]}")
            logging.debug(f"【デバッグログ】parent_text 全体の長さ: {len(parent_text)}")
        else:
            logging.debug(f"【デバッグログ】parent_text.strip() が False と評価されました。parent_textの先頭50文字: {parent_text[:50]}")
            logging.debug(f"【デバッグログ】parent_text 全体の長さ: {len(parent_text)}")
